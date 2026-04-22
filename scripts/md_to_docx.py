#!/usr/bin/env python3
"""Convert the Bitewise vector-DB plan markdown into a Word document."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor


BRAND_BLUE = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0x2E, 0x75, 0xB6)
GRAY = RGBColor(0x59, 0x59, 0x59)
CODE_BG = "F2F2F2"
TABLE_HEADER_BG = "1F4E79"
ZEBRA_BG = "F8F9FB"


def set_cell_bg(cell, color_hex: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def set_paragraph_bg(paragraph, color_hex: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    pPr.append(shd)


def add_inline_runs(paragraph, text: str, base_color=None, base_bold=False) -> None:
    """Render inline markdown (bold, code, italic) into runs."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.font.size = Pt(10.5)
            if base_color:
                run.font.color.rgb = base_color
            if base_bold:
                run.bold = True
        token = match.group()
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.size = Pt(10.5)
            if base_color:
                run.font.color.rgb = base_color
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xB0, 0x3A, 0x2E)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
            run.font.size = Pt(10.5)
            if base_color:
                run.font.color.rgb = base_color
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(10.5)
        if base_color:
            run.font.color.rgb = base_color
        if base_bold:
            run.bold = True


def add_heading(doc: Document, text: str, level: int) -> None:
    text = text.strip()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
    p.paragraph_format.space_after = Pt(6)
    sizes = {1: 22, 2: 17, 3: 14, 4: 12}
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = BRAND_BLUE if level <= 2 else ACCENT
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_code_block(doc: Document, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, CODE_BG)
    cell.text = ""
    for idx, raw in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(raw)
        run.font.name = "Menlo"
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_quote(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_bg(p, "FFF7E6")
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRAY


def add_bullet(doc: Document, text: str, level: int = 0, numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.6)
    p.paragraph_format.space_after = Pt(2)
    add_inline_runs(p, text)


def build_table(doc: Document, header, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        set_cell_bg(cell, TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(cell_text.strip())
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= len(header):
                continue
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            if r_idx % 2 == 1:
                set_cell_bg(cell, ZEBRA_BG)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline_runs(p, cell_text.strip())
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def parse_and_render(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Helvetica Neue"
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn("w:eastAsia"), "PingFang SC")
    style.font.size = Pt(10.5)

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_buffer = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buffer)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        stripped = line.rstrip()

        if not stripped.strip():
            doc.add_paragraph()
            i += 1
            continue

        if stripped.strip() == "---":
            add_divider(doc)
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            add_heading(doc, heading_match.group(2), level)
            i += 1
            continue

        if stripped.startswith(">"):
            add_quote(doc, stripped.lstrip(">").strip())
            i += 1
            continue

        if stripped.lstrip().startswith("|") and "|" in stripped[1:]:
            header_line = stripped.lstrip()
            separator = lines[i + 1].strip() if i + 1 < len(lines) else ""
            if re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?$", separator):
                header_cells = [c.strip() for c in header_line.strip("|").split("|")]
                data_rows = []
                j = i + 2
                while j < len(lines) and lines[j].lstrip().startswith("|"):
                    row_cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                    data_rows.append(row_cells)
                    j += 1
                build_table(doc, header_cells, data_rows)
                i = j
                continue

        list_match = re.match(r"^(\s*)([-*])\s+(.*)$", stripped)
        if list_match:
            indent = len(list_match.group(1)) // 2
            add_bullet(doc, list_match.group(3), level=indent, numbered=False)
            i += 1
            continue

        num_match = re.match(r"^(\s*)(\d+)\.\s+(.*)$", stripped)
        if num_match:
            indent = len(num_match.group(1)) // 2
            add_bullet(doc, num_match.group(3), level=indent, numbered=True)
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_inline_runs(p, stripped)
        i += 1

    doc.save(str(docx_path))


if __name__ == "__main__":
    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    parse_and_render(src, dst)
    print(f"Generated: {dst}")
