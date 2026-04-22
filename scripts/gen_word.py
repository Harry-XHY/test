"""Generate SocialChef Word document on Desktop."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SocialChef")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("海外社媒营销自动化平台\n市场调研 · 可行性分析 · 产品规划")
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_p.add_run("2026 年 4 月 13 日")
run3.font.size = Pt(14)
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# TOC
# ══════════════════════════════════════════════════════════
doc.add_heading("目录", level=1)
toc_items = [
    "一、项目概述",
    "二、市场调研",
    "  2.1 市场规模与趋势",
    "  2.2 竞品分析",
    "  2.3 市场空白与机会",
    "  2.4 定价策略建议",
    "三、技术可行性分析",
    "  3.1 竞品监控",
    "  3.2 视频转录",
    "  3.3 AI 脚本生成",
    "  3.4 AI 视频生成",
    "  3.5 多平台发布",
    "  3.6 综合成本估算",
    "四、产品需求文档 (PRD)",
    "  4.1 产品愿景与目标",
    "  4.2 用户画像",
    "  4.3 功能模块",
    "  4.4 非功能需求",
    "五、系统架构与技术栈",
    "六、实施路线图",
    "  6.1 Phase 1: MVP（8周）",
    "  6.2 Phase 2: 扩展（6周）",
    "  6.3 Phase 3: 规模化（8周）",
    "  6.4 Phase 4: SaaS 产品化",
    "七、风险分析与应对",
    "八、MVP 验证标准",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 一、项目概述
# ══════════════════════════════════════════════════════════
doc.add_heading("一、项目概述", level=1)

doc.add_paragraph(
    "SocialChef 是一个全新的独立项目，目标是构建一个通用的海外社媒营销自动化平台，"
    "服务于在国内运营海外社交媒体账号的企业和团队。"
)

doc.add_heading("核心理念", level=2)
doc.add_paragraph(
    "通过 AI 全自动化营销流水线，将内容生产成本压到极致。"
    "一个运营人员即可管理 100+ 个社媒账号。"
    "以极低价格或免费提供社媒托管服务，快速签约大量客户。"
)

doc.add_heading("核心工作流", level=2)
doc.add_paragraph(
    "竞品监测 → 爆款识别 → 视频转文字 → AI 生成脚本 → "
    "AI 生成视频 (×3 变体) → 人工审核 → 多平台自动发布"
)

doc.add_heading("商业模式", level=2)
bullets = [
    "AI 自动化将单客户运营成本压至 < $50/月",
    "对客户免费或收极低费用（快速获客）",
    "用「社媒营销免费托管」作为获客手段，快速扩大客户网络",
    "长期可独立成为通用 Marketing Automation SaaS 产品",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 二、市场调研
# ══════════════════════════════════════════════════════════
doc.add_heading("二、市场调研", level=1)

doc.add_heading("2.1 市场规模与趋势", level=2)
doc.add_paragraph("社交媒体营销市场正在经历爆发式增长，关键数据如下：")
stats = [
    "2025 年全球社交媒体广告支出达 2,767 亿美元",
    "74% 的消费者通过社交媒体决定消费选择（以餐饮为例）",
    "55% 的企业已开始使用 AI 创建营销内容",
    "TikTok 短视频营销采用率从 2023 年的 26% 翻倍至 2024 年的 48%",
    "企业外包社媒管理费通常为 $500-$1,500/月",
    "90% 的企业认为社交媒体对数字营销「非常/极其重要」",
]
for s in stats:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("2.2 竞品分析", level=2)

doc.add_heading("社交媒体管理平台（通用型）", level=3)
add_table(
    ["平台", "定价", "核心功能", "局限性"],
    [
        ["Hootsuite", "$99/用户/月起", "多平台管理、排期发布、分析报表", "价格偏高；不生成内容；无视频制作"],
        ["Sprout Social", "$199/席位/月起", "深度分析、竞品基准、社交聆听", "价格极高；面向企业；无内容创作"],
        ["Buffer", "免费起步，$6/月/渠道", "简洁易用、11 平台支持、AI 助手", "功能简单；无视频制作；无竞品监控"],
        ["Later", "$18.75/月起", "视觉内容规划、AI 排期、Canva 集成", "偏向图片；视频能力弱"],
        ["Loomly", "$26/月起", "团队协作审批、内容日历", "无 AI 内容生成；无视频制作"],
        ["SocialBee", "$29/月起", "AI Copilot、常青内容循环", "无视频制作；不监控竞品视频"],
    ],
)

doc.add_heading("AI 视频/内容工具", level=3)
add_table(
    ["平台", "定价", "核心功能", "局限性"],
    [
        ["Revid.ai", "免费起步", "AI 脚本→AI 视频→一键发布", "无竞品监控；无行业定制；模板拼接"],
        ["OpusClip", "$15/月起", "AI 从长视频提取精华片段", "仅剪辑，不能从零生成"],
        ["Vizard AI", "$14.5/月起", "视频裁剪、多平台格式化", "仅剪辑工具"],
        ["Pictory", "$23/月起", "文字转视频、自动字幕", "模板式视频；画面质量一般"],
        ["Predis.ai", "免费起步", "AI 文案 + 模板视频 + 竞品洞察", "模板式视频，非 AI 原创"],
    ],
)

doc.add_heading("AI 视频生成引擎（底层能力）", level=3)
add_table(
    ["平台", "特点", "适用场景"],
    [
        ["HeyGen", "AI 数字人/虚拟主播、40+ 语言", "口播类视频、产品介绍"],
        ["Seedance 2.0（字节）", "文字/图片→视频、1080p、电影级画质", "场景类视频、产品展示"],
        ["Kling AI（可灵）", "文字→视频、45-60 秒生成", "通用视频生成"],
    ],
)

doc.add_heading("餐饮/垂直自动化工具", level=3)
add_table(
    ["平台", "定价", "核心功能", "局限性"],
    [
        ["Glow Social", "$49/月", "全自动：12 条/月，自动发布多平台", "仅图文帖子，无视频"],
        ["Platr.ai", "未公开", "每周 2-3 条定制帖子", "发帖量少；无视频"],
        ["Bloom Intelligence", "未公开", "CDP + 自动化 + 声誉管理", "偏 CRM，非内容创作"],
    ],
)

doc.add_heading("2.3 市场空白与机会", level=2)

doc.add_paragraph(
    "没有任何一个现有工具能完成「监控→识别→转录→脚本→视频→审核→发布」的完整闭环。",
)

doc.add_heading("全流程覆盖对比", level=3)
add_table(
    ["流程环节", "Hootsuite", "Revid.ai", "OpusClip", "Glow Social", "SocialChef"],
    [
        ["监控竞品视频", "部分", "—", "—", "—", "全覆盖"],
        ["识别爆款内容", "—", "部分", "病毒性评分", "—", "全覆盖"],
        ["视频→文字转录", "—", "—", "—", "—", "全覆盖"],
        ["AI 生成脚本", "—", "有", "—", "—", "全覆盖"],
        ["AI 生成视频", "—", "模板式", "仅剪辑", "—", "全覆盖 (HeyGen/Kling)"],
        ["多变体 A/B", "—", "—", "—", "—", "全覆盖 (~3 变体)"],
        ["人工审核", "有", "—", "—", "有", "全覆盖"],
        ["多平台发布", "有", "有", "—", "有", "全覆盖"],
    ],
)

doc.add_heading("核心差异化", level=3)
diffs = [
    "全流程自动化 — 业内首个完整闭环，用户几乎零操作",
    "AI 原生视频生成 — 基于 HeyGen + Kling 的新一代 AI 视频，非模板拼接",
    "竞品智能学习 — 自动监控同行爆款内容，提取成功要素",
    "多变体 + 人工选优 — 每次生成 ~3 个变体，保留人工审美判断",
    "极低单客成本 — AI 全自动化使边际成本极低，远低于人工外包",
]
for d in diffs:
    doc.add_paragraph(d, style="List Bullet")

doc.add_heading("2.4 定价策略建议", level=2)
add_table(
    ["方案", "建议定价", "对标", "价值主张"],
    [
        ["基础版", "$49-79/月", "Glow Social ($49)", "每月 8-12 条视频内容"],
        ["标准版", "$129-199/月", "人工外包 ($500-1500) 的 1/5", "每月 20+ 条，含竞品监控 + 多变体"],
        ["连锁/企业版", "$299+/月", "Sprout Social ($199+)", "多账号管理、统一品牌、本地化"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 三、技术可行性分析
# ══════════════════════════════════════════════════════════
doc.add_heading("三、技术可行性分析", level=1)

doc.add_heading("3.1 竞品监控", level=2)
doc.add_paragraph("技术复杂度：★★★★ (难)")
add_table(
    ["平台", "官方 API", "能力", "限制", "推荐方案"],
    [
        ["YouTube", "Data API v3", "可获取任何公开视频/频道详细数据", "每天 10,000 单位配额", "直接用官方 API（免费）"],
        ["TikTok", "Research API / Display API", "需审批或用户授权", "无法直接监控竞品", "Apify TikTok Scraper"],
        ["Instagram", "Graph API", "仅自有账号数据", "无法获取竞品指标", "Apify / Ensemble Data"],
        ["Facebook", "Pages API", "需 Meta 审核", "审批困难", "第三方服务"],
    ],
)

doc.add_heading("第三方数据服务", level=3)
add_table(
    ["服务", "覆盖平台", "价格", "备注"],
    [
        ["Apify", "TikTok/IG/YouTube", "$49/月起", "现成 scraper，按用量计费"],
        ["Ensemble Data", "TikTok/IG/YouTube", "$200-500/月", "无需授权，直接抓取公开数据"],
        ["RapidAPI 各类 scraper", "各平台", "$10-100/月", "稳定性不保证"],
    ],
)

doc.add_heading("法律/TOS 风险", level=3)
legal = [
    "TikTok/Instagram TOS 明确禁止未授权的 scraping",
    "YouTube 官方 API 提供了充足的公开数据访问能力，合法合规",
    "通过第三方服务商（Apify, Ensemble Data）间接获取数据可将法律风险转嫁给服务商",
    "实际操作中大量第三方工具在正常运营",
]
for l in legal:
    doc.add_paragraph(l, style="List Bullet")

doc.add_heading("3.2 视频转录", level=2)
doc.add_paragraph("技术复杂度：★ (简单)")

doc.add_paragraph("视频下载：yt-dlp（开源免费，支持 1000+ 网站）")
add_table(
    ["服务", "价格", "准确度", "多语言", "特色"],
    [
        ["OpenAI Whisper API", "$0.006/分钟", "优秀", "99+ 语言", "性价比最高（推荐）"],
        ["Deepgram", "$0.0043/分钟", "优秀", "36+ 语言", "实时转录能力强"],
        ["AssemblyAI", "$0.006/分钟", "优秀", "多语言", "内容审核等附加功能"],
    ],
)
doc.add_paragraph("成本估算（15 视频/月，每个 60 秒）：$0.09/月/客户 — 几乎可忽略")

doc.add_heading("3.3 AI 脚本生成", level=2)
doc.add_paragraph("技术复杂度：★★ (简单)")
add_table(
    ["服务", "模型", "每个脚本预估成本"],
    [
        ["Claude API", "Sonnet 4", "$0.01-0.03/脚本"],
        ["Claude API", "Haiku", "$0.001-0.005/脚本"],
        ["OpenAI", "GPT-4o", "$0.01-0.02/脚本"],
        ["OpenAI", "GPT-4o-mini", "$0.001-0.003/脚本"],
    ],
)
doc.add_paragraph("成本估算（15 脚本/月）：$0.15-0.45/月/客户")

doc.add_heading("3.4 AI 视频生成", level=2)
doc.add_paragraph("技术复杂度：★★★★★ (最难) — 整个流水线中技术最复杂、成本最高的环节")
add_table(
    ["服务", "API", "价格", "生成时间", "质量"],
    [
        ["HeyGen", "有 REST API", "$48-120/月 + 超出 $1-3/分钟", "2-10 分钟", "高（真人数字人）"],
        ["Kling AI（可灵）", "有 API", "~¥0.14-0.28/秒", "2-5 分钟", "高（视觉效果好）"],
        ["Runway Gen-3/4", "有 API", "$0.05-0.25/秒", "1-5 分钟", "高"],
        ["D-ID", "有 REST API", "$0.03-0.10/秒", "1-5 分钟", "中高（数字人）"],
        ["Seedance（字节）", "API 内测中", "待定", "未知", "高"],
    ],
)

doc.add_heading("30 秒视频成本估算", level=3)
add_table(
    ["方案", "成本/视频", "质量"],
    [
        ["纯 Kling AI", "$0.60-1.20", "高"],
        ["纯 Runway", "$1.50-7.50", "高"],
        ["HeyGen 数字人 + Kling 混合", "$2-5", "最高"],
        ["Synthesia", "$22-67", "企业级（不推荐大量）"],
    ],
)
doc.add_paragraph("每月成本（15 视频 × 3 变体 = 45 视频）：$27-225/月/客户")

doc.add_heading("3.5 多平台发布", level=2)
doc.add_paragraph("技术复杂度：★★★ (中等)")
add_table(
    ["平台", "官方 API", "要求"],
    [
        ["TikTok", "Content Posting API", "需开发者账号审批 + OAuth"],
        ["Instagram", "Content Publishing API", "Business 账号 + Meta 审核"],
        ["YouTube", "Data API v3 (upload)", "Google OAuth + 审核"],
        ["Facebook", "Pages API", "需 Meta app review"],
    ],
)

doc.add_heading("推荐：第三方发布服务", level=3)
add_table(
    ["服务", "平台覆盖", "价格", "备注"],
    [
        ["Ayrshare（推荐）", "TikTok/IG/YouTube/FB", "$25-250/月", "API 设计好，支持视频"],
        ["Publer", "TikTok/IG/YouTube/FB", "$12-42/月", "价格友好"],
        ["Buffer", "IG/YouTube/FB/TikTok", "$6-120/月", "成熟稳定"],
    ],
)

doc.add_heading("3.6 综合成本估算", level=2)
doc.add_paragraph("每客户每月成本（15 视频/月，每视频 3 变体）：")
add_table(
    ["项目", "低成本方案", "标准方案", "备注"],
    [
        ["竞品监控", "$50 (Apify)", "$200 (Ensemble Data)", "多客户可共享"],
        ["视频下载 + 转录", "$1", "$2", "Whisper API"],
        ["AI 脚本生成", "$0.50", "$2", "Claude/GPT"],
        ["AI 视频生成", "$30", "$150", "Kling vs 混合方案"],
        ["多平台发布", "$25 (Ayrshare)", "$25", "多客户可共享"],
        ["合计", "~$107/月", "~$379/月", ""],
    ],
)
doc.add_paragraph(
    "规模化后（100 个客户）：竞品监控和发布工具成本被分摊，"
    "每客户有效成本降至 ~$32-155/月。视频生成是主要成本（占 60-70%）。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 四、PRD
# ══════════════════════════════════════════════════════════
doc.add_heading("四、产品需求文档 (PRD)", level=1)

doc.add_heading("4.1 产品愿景与目标", level=2)
doc.add_paragraph("愿景：成为全球领先的 AI 驱动社媒营销自动化平台。")
doc.add_paragraph("核心目标：")
kpis = [
    "3 个月内签约 50 个客户",
    "每个客户每周产出 3+ 条视频内容",
    "单人运营可管理 100+ 客户",
    "每客户月运营成本 < $50（规模化后）",
]
for k in kpis:
    doc.add_paragraph(k, style="List Bullet")

doc.add_heading("4.2 用户画像", level=2)

doc.add_heading("内部运营人员", level=3)
doc.add_paragraph(
    "团队规模 1-2 人，负责管理所有客户的社媒营销。"
    "核心需求：效率最大化、批量操作、一键审核发布。"
)

doc.add_heading("客户（企业/品牌）", level=3)
doc.add_paragraph(
    "需要海外社媒营销但缺乏团队/技能的企业。"
    "核心需求：低成本、高质量内容、稳定更新频率。"
)

doc.add_heading("4.3 功能模块", level=2)

modules = [
    ("竞品监测模块", [
        ("P0", "手动添加竞品账号 URL"),
        ("P0", "自动抓取竞品视频列表与基础指标"),
        ("P0", "爆款识别算法（播放量/互动率/发布时间综合评分）"),
        ("P1", "定时自动抓取（每日/每周）"),
        ("P1", "爆款推送通知（邮件/webhook）"),
        ("P2", "趋势分析报告（周报/月报）"),
    ]),
    ("内容转录模块", [
        ("P0", "视频下载（yt-dlp）"),
        ("P0", "语音转文字（Whisper API）"),
        ("P0", "多语言支持（泰语/英语/中文）"),
        ("P1", "关键词提取与标签生成"),
    ]),
    ("AI 脚本生成模块", [
        ("P0", "输入爆款脚本 + 客户资料 → 生成新脚本"),
        ("P0", "支持多时长变体（15s/30s/60s）"),
        ("P0", "每次生成 3 个脚本变体"),
        ("P1", "脚本编辑器（运营可手动调整）"),
        ("P1", "品牌调性配置（风格/语气/关键词）"),
        ("P2", "历史表现反馈优化 prompt"),
    ]),
    ("AI 视频生成模块", [
        ("P0", "脚本 → HeyGen 数字人视频"),
        ("P0", "FFmpeg 拼接（数字人 + 素材 + 字幕 + 配乐）"),
        ("P1", "接入 Kling AI 生成场景视频"),
        ("P1", "每脚本 3 个视频变体"),
        ("P2", "自动选择最佳 AI 引擎"),
    ]),
    ("审核与发布模块", [
        ("P0", "视频预览界面"),
        ("P0", "一键通过/拒绝"),
        ("P0", "Ayrshare API 发布到 TikTok"),
        ("P1", "排期发布（设定时间）"),
        ("P1", "扩展到 Instagram/YouTube/Facebook"),
        ("P2", "A/B 测试（发布多变体对比效果）"),
    ]),
    ("客户管理模块", [
        ("P0", "客户 CRUD（名称、产品、品牌调性、社媒账号）"),
        ("P0", "每个客户独立的竞品监控列表"),
        ("P0", "发布历史记录"),
        ("P1", "数据看板（播放量/互动/增长趋势）"),
        ("P2", "客户自助查看报告"),
    ]),
]

for mod_name, features in modules:
    doc.add_heading(mod_name, level=3)
    add_table(
        ["优先级", "功能描述"],
        [[p, desc] for p, desc in features],
    )

doc.add_heading("4.4 非功能需求", level=2)
nfrs = [
    "性能：视频生成任务异步处理，不阻塞 UI；支持 100+ 客户并发",
    "安全：API Key 加密存储；客户数据隔离；HTTPS 全链路加密",
    "合规：遵守各平台 TOS；PDPA（泰国）/ GDPR 数据保护合规",
    "可用性：99.5% uptime；任务失败自动重试",
    "可扩展：模块化架构，便于接入新的 AI 引擎和社交平台",
]
for n in nfrs:
    doc.add_paragraph(n, style="List Bullet")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 五、技术栈
# ══════════════════════════════════════════════════════════
doc.add_heading("五、系统架构与技术栈", level=1)

add_table(
    ["层", "技术", "说明"],
    [
        ["前端", "Next.js 16 + Tailwind CSS + shadcn/ui", "响应式管理后台"],
        ["后端", "Python FastAPI + Celery + Redis", "API 服务 + 异步任务队列"],
        ["数据库", "PostgreSQL (Supabase)", "客户数据、任务记录"],
        ["AI 脚本", "Claude API / OpenAI API", "营销脚本生成"],
        ["AI 转录", "OpenAI Whisper API", "视频语音转文字"],
        ["AI 视频", "HeyGen API + Kling API + FFmpeg", "数字人 + 场景视频 + 拼接"],
        ["竞品监控", "YouTube Data API v3 + Apify", "公开数据抓取"],
        ["发布", "Ayrshare API", "多平台统一发布"],
        ["部署", "Vercel (前端) + Railway (后端)", "自动化 CI/CD"],
    ],
)

doc.add_heading("架构图", level=2)
arch = (
    "┌─────────────┐     ┌──────────────┐     ┌─────────────┐\n"
    "│  Next.js UI  │────▶│  FastAPI 后端  │────▶│  PostgreSQL  │\n"
    "│  (Vercel)    │     │  (Railway)    │     │  (Supabase)  │\n"
    "└─────────────┘     └──────┬───────┘     └─────────────┘\n"
    "                           │\n"
    "                    ┌──────┴───────┐\n"
    "                    │  Celery 队列  │\n"
    "                    │  (Redis)     │\n"
    "                    └──────┬───────┘\n"
    "          ┌────────┬───────┼────────┬──────────┐\n"
    "          ▼        ▼       ▼        ▼          ▼\n"
    "      Apify    Whisper   Claude   HeyGen   Ayrshare\n"
    "     (监控)    (转录)    (脚本)   (视频)    (发布)\n"
)
p = doc.add_paragraph(arch)
for run in p.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(9)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 六、路线图
# ══════════════════════════════════════════════════════════
doc.add_heading("六、实施路线图", level=1)

doc.add_heading("6.1 Phase 1: MVP（第 1-8 周）", level=2)
doc.add_paragraph("目标：5 个种子客户，仅 TikTok，验证核心价值链")
mvp_tasks = [
    "Task 1: 项目脚手架搭建 — 全新独立仓库，前后端基础框架",
    "Task 2: 竞品监控模块（半自动）— 手动输入账号，自动抓取 + 爆款识别",
    "Task 3: 视频转录模块 — yt-dlp + Whisper API",
    "Task 4: AI 脚本生成模块 — Claude API + Prompt 模板",
    "Task 5: AI 视频生成模块（半自动）— HeyGen 数字人 + FFmpeg 拼接",
    "Task 6: 审核与发布模块 — 预览/审核 UI + Ayrshare 发布",
    "Task 7: 客户管理 — 多客户后台",
]
for t in mvp_tasks:
    doc.add_paragraph(t, style="List Bullet")

doc.add_heading("6.2 Phase 2: 扩展（第 9-14 周）", level=2)
doc.add_paragraph("目标：50 个客户，多平台支持")
p2 = [
    "接入 Kling AI，支持更丰富的 AI 视频生成",
    "每脚本 3 变体视频",
    "扩展到 Instagram Reels + YouTube Shorts",
    "竞品监控自动化：定时抓取 + 爆款自动推送通知",
    "效果追踪：发布后播放量/互动数据回收",
]
for item in p2:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("6.3 Phase 3: 规模化（第 15-22 周）", level=2)
doc.add_paragraph("目标：200 个客户，全自动化")
p3 = [
    "全自动流水线：监测到爆款 → 自动生成脚本 + 视频 → 推送审核",
    "智能内容策略：基于历史数据推荐最佳发布时间/内容类型",
    "A/B 测试闭环：自动对比变体效果，优化生成策略",
    "多语言支持：泰语/英语/中文",
]
for item in p3:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("6.4 Phase 4: SaaS 产品化（第 23 周+）", level=2)
p4 = [
    "自助注册 + 订阅付费",
    "白标方案（给其他营销公司用）",
    "API 开放",
]
for item in p4:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 七、风险
# ══════════════════════════════════════════════════════════
doc.add_heading("七、风险分析与应对", level=1)

add_table(
    ["风险", "影响", "概率", "应对措施"],
    [
        ["AI 视频质量不达标", "高", "中", "MVP 阶段半自动（人工 + AI 混合）；持续跟进技术迭代"],
        ["TikTok/IG 数据获取受限", "高", "中", "通过第三方服务商间接获取；YouTube 作为兜底"],
        ["平台对 AI 内容降权", "中", "低-中", "保持人工审核环节；混合真实素材 + AI 生成"],
        ["第三方 API 价格变动", "中", "中", "架构支持多引擎切换；关注新兴替代方案"],
        ["内容同质化", "中", "中", "品牌调性配置差异化；多变体策略"],
        ["合规风险（PDPA/GDPR）", "高", "低", "数据最小化原则；用户授权机制"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 八、验证标准
# ══════════════════════════════════════════════════════════
doc.add_heading("八、MVP 验证标准", level=1)

checks = [
    "5 个种子客户成功接入系统",
    "每周每个客户产出 3+ 条 TikTok 视频",
    "单人可管理全部 5 个客户的社媒运营",
    "从爆款发现到视频发布 < 2 小时（含人工审核）",
    "每客户月运营成本 < $50",
]
for c in checks:
    doc.add_paragraph("☐  " + c)

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("— SocialChef 项目规划文档 · 2026.04.13 —")
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(10)

# Save
output = os.path.expanduser("~/Desktop/SocialChef_项目规划.docx")
doc.save(output)
print(f"Saved to: {output}")
