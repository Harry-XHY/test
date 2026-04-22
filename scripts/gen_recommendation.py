"""Generate build-vs-buy recommendation Word document."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)


def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


# ══════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SocialChef")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run("Build vs Buy 决策分析\n自建平台 vs 工具组合方案对比")
run2.font.size = Pt(18)
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = date_p.add_run("2026 年 4 月 13 日")
run3.font.size = Pt(14)
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 核心结论
# ══════════════════════════════════════════════════════════
doc.add_heading("核心结论", level=1)

p = doc.add_paragraph()
run = p.add_run("不推荐一上来就自己开发一套完整平台。")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

doc.add_paragraph(
    "建议分两步走：先用现有工具组合（n8n + 各 API）快速验证商业模式，"
    "确认客户需求和 AI 视频质量后，再启动自建平台。"
)

doc.add_paragraph("核心逻辑：先花 $128/月验证，别花 $10,000+ 赌一个假设。")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 两个方案对比
# ══════════════════════════════════════════════════════════
doc.add_heading("一、两个方案全面对比", level=1)

add_table(
    ["维度", "方案 A：工具组合", "方案 B：自建平台"],
    [
        ["上线时间", "1-2 周", "8-12 周"],
        ["前期投入", "~$0（仅 API 月费）", "$10,000-30,000（开发成本）"],
        ["月运营成本", "~$128/月", "~$80/月（仅 API 费） + 服务器"],
        ["可管理客户数", "5-30 个", "100+ 个"],
        ["灵活性", "中等（受工具限制）", "高（完全定制）"],
        ["维护成本", "低", "高（需持续开发维护）"],
        ["SaaS 化可能", "不可能", "可以"],
        ["技术风险", "低（成熟工具）", "高（AI API 变化快）"],
        ["验证速度", "快（1-2 周出结果）", "慢（2-3 月才能验证）"],
        ["退出成本", "几乎为零", "沉没成本高"],
    ],
)

# ══════════════════════════════════════════════════════════
# 方案 A 详细
# ══════════════════════════════════════════════════════════
doc.add_heading("二、方案 A：工具组合（推荐先做）", level=1)

doc.add_heading("2.1 整体架构", level=2)

arch = (
    "                     n8n 工作流引擎（核心调度）\n"
    "                            |\n"
    "     +----------+----------+----------+----------+\n"
    "     |          |          |          |          |\n"
    "   Apify     Whisper    Claude     HeyGen    Ayrshare\n"
    "  (竞品监控)  (转录)    (脚本)     (视频)     (发布)\n"
    "     |          |          |          |          |\n"
    "     +----------+----------+----------+----------+\n"
    "                            |\n"
    "              飞书/Notion 看板（人工审核）\n"
)
p = doc.add_paragraph(arch)
for run in p.runs:
    run.font.name = "Courier New"
    run.font.size = Pt(9)

doc.add_heading("2.2 工具清单与费用", level=2)

add_table(
    ["工具", "月费", "作用", "替代方案"],
    [
        ["n8n（自托管）", "免费", "工作流引擎，串联所有 API", "Make.com ($9/月)"],
        ["Apify", "$49/月", "TikTok/IG/YouTube 数据抓取", "Ensemble Data ($200/月)"],
        ["OpenAI Whisper API", "~$1/月", "视频语音转文字", "Deepgram ($4/月)"],
        ["Claude API", "~$5/月", "AI 营销脚本生成", "GPT-4o ($3/月)"],
        ["HeyGen", "$48/月", "AI 数字人视频生成", "D-ID ($30/月)"],
        ["Ayrshare", "$25/月", "多平台一键发布", "Publer ($12/月)"],
        ["合计", "~$128/月", "", "最低可压至 ~$55/月"],
    ],
)

doc.add_heading("2.3 工作流详细设计", level=2)

doc.add_heading("流程 1：竞品监控 + 爆款识别", level=3)
steps1 = [
    "运营人员在 n8n 表单中输入竞品账号 URL",
    "n8n 定时触发（每天/每周）Apify 抓取竞品最新视频",
    "n8n 内置 JavaScript 节点计算爆款评分（播放量 x 互动率 / 发布天数）",
    "评分超过阈值的视频自动推送到飞书/Slack 通知",
]
for i, s in enumerate(steps1, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("流程 2：转录 + 脚本生成", level=3)
steps2 = [
    "运营从通知中选择要分析的爆款视频",
    "n8n 调用 yt-dlp（通过 Execute Command 节点）下载视频音频",
    "音频发送至 Whisper API 转录为文字",
    "转录文字 + 客户资料 → 发送至 Claude API",
    "Claude 返回 3 个脚本变体（15s/30s/60s）",
    "脚本推送到飞书/Notion 看板供审核",
]
for i, s in enumerate(steps2, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("流程 3：视频生成", level=3)
steps3 = [
    "运营审核通过脚本后，在看板标记「通过」",
    "n8n 监听看板状态变更，自动触发 HeyGen API",
    "HeyGen 生成数字人口播视频（支持多语言、多风格）",
    "视频生成完成后推送到审核看板",
]
for i, s in enumerate(steps3, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("流程 4：审核 + 发布", level=3)
steps4 = [
    "运营在看板预览视频，标记「发布」或「拒绝」",
    "标记「发布」后，n8n 自动调用 Ayrshare API",
    "Ayrshare 将视频发布到 TikTok / Instagram / YouTube",
    "发布结果回写到看板记录",
]
for i, s in enumerate(steps4, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("2.4 每客户运营成本", level=2)
doc.add_paragraph("假设每月 15 个视频，每视频 1 个变体（MVP 阶段）：")
add_table(
    ["项目", "成本", "说明"],
    [
        ["竞品监控", "$0.50/客户", "Apify $49 分摊 100 个客户"],
        ["转录", "$0.09/客户", "15 分钟 x $0.006"],
        ["脚本生成", "$0.45/客户", "15 个脚本 x $0.03"],
        ["视频生成", "$30-48/客户", "HeyGen 按量或包月"],
        ["发布", "$0.25/客户", "Ayrshare $25 分摊"],
        ["合计", "~$31-49/客户/月", "视频生成是主要成本"],
    ],
)

doc.add_heading("2.5 方案 A 的局限性", level=2)
limits = [
    "没有统一的管理界面 — 需要在 n8n、飞书、各 API 后台之间切换",
    "客户数量受限 — 超过 30 个客户后，看板管理会变得混乱",
    "无法 SaaS 化 — 不能给客户提供自助界面",
    "视频只支持数字人口播 — 无法做复杂的多镜头拼接",
    "依赖工具稳定性 — 如果 Apify 或 HeyGen 出问题，整个流程中断",
]
for l in limits:
    doc.add_paragraph(l, style="List Bullet")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 方案 B 详细
# ══════════════════════════════════════════════════════════
doc.add_heading("三、方案 B：自建平台（验证后再做）", level=1)

doc.add_heading("3.1 什么时候该启动自建？", level=2)
doc.add_paragraph("当你遇到以下信号时，说明该自建了：")

signals = [
    "客户数超过 30-50 个，工具组合管理不过来",
    "核心流程已稳定运行 2-3 个月，知道哪些环节需要定制",
    "有明确的付费客户或营收数据，商业模式验证成功",
    "需要给客户提供自助界面（SaaS 化）",
    "AI 视频质量已确认满足客户需求",
]
for s in signals:
    doc.add_paragraph(s, style="List Bullet")

doc.add_heading("3.2 自建的优势", level=2)
advantages = [
    "统一管理后台 — 一个界面管理所有客户、所有流程",
    "深度定制 — 根据实际需求定制爆款算法、脚本模板、视频风格",
    "规模化能力 — 支持 100+ 客户并发，自动化程度更高",
    "SaaS 化 — 可以开放给其他营销团队使用，形成新的营收来源",
    "数据闭环 — 发布效果数据反馈到脚本生成，持续优化",
    "多引擎切换 — 根据价格/质量自动选择最优 AI 引擎",
]
for a in advantages:
    doc.add_paragraph(a, style="List Bullet")

doc.add_heading("3.3 自建的成本估算", level=2)
add_table(
    ["项目", "估算", "说明"],
    [
        ["前端开发", "$3,000-5,000", "Next.js 管理后台"],
        ["后端开发", "$4,000-8,000", "FastAPI + Celery + 所有 API 集成"],
        ["测试 + 部署", "$1,000-2,000", "E2E 测试、CI/CD"],
        ["月维护成本", "$500-1,000/月", "Bug 修复、API 变更适配"],
        ["月基础设施", "$50-100/月", "Vercel + Railway + Supabase"],
        ["合计（首年）", "$14,000-27,000", "含开发 + 12 个月运维"],
    ],
)

doc.add_heading("3.4 自建的风险", level=2)
risks = [
    "开发期间无收入 — 8-12 周开发期 = 2-3 个月空窗期",
    "AI API 变化快 — HeyGen/Kling 的 API 和定价可能随时变",
    "需求可能变 — 验证前的假设可能与实际需求不符",
    "维护负担 — 需要持续投入开发资源",
]
for r in risks:
    doc.add_paragraph(r, style="List Bullet")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 推荐路径
# ══════════════════════════════════════════════════════════
doc.add_heading("四、推荐执行路径", level=1)

doc.add_heading("第一阶段：快速验证（第 1-4 周）", level=2)
p = doc.add_paragraph()
run = p.add_run("方案 A：工具组合")
run.bold = True
phase1 = [
    "第 1 周：搭建 n8n 工作流，跑通完整链路（监控→转录→脚本→视频→发布）",
    "第 2 周：找 3-5 个种子客户，开始试运营",
    "第 3-4 周：收集反馈，优化 prompt、调整视频风格、验证客户满意度",
]
for item in phase1:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("验证指标", level=3)
add_table(
    ["指标", "目标值", "判断标准"],
    [
        ["客户满意度", "3/5 个客户续约意向", "客户愿意继续使用"],
        ["视频质量", "70%+ 视频通过人工审核", "AI 生成质量可接受"],
        ["运营效率", "1 人管理 5 个客户", "流程可运转"],
        ["获客成本", "每客户 < $200 获客", "商业模式可行"],
        ["内容效果", "平均播放量 > 行业中位数", "内容有竞争力"],
    ],
)

doc.add_heading("第二阶段：扩大验证（第 5-12 周）", level=2)
p = doc.add_paragraph()
run = p.add_run("继续方案 A，优化工作流")
run.bold = True
phase2 = [
    "扩大到 10-20 个客户",
    "优化爆款识别算法（基于实际数据）",
    "测试不同类型的脚本和视频风格",
    "记录所有流程中的痛点和瓶颈",
    "开始评估自建平台的需求优先级",
]
for item in phase2:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("第三阶段：决策点（第 12 周）", level=2)
p = doc.add_paragraph()
run = p.add_run("Go / No-Go 决策")
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

doc.add_paragraph("如果验证成功（大部分指标达标）：")
go = [
    "启动方案 B 自建平台",
    "基于 12 周的真实运营数据设计功能",
    "优先开发工具组合中的痛点功能",
    "保持工具组合继续运营，自建平台并行开发",
]
for item in go:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph("如果验证失败（多数指标未达标）：")
nogo = [
    "分析失败原因（是需求问题还是执行问题？）",
    "调整方向或止损退出",
    "总投入仅 ~$500-1,000（工具费用），损失可控",
]
for item in nogo:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("第四阶段：自建平台（第 13-24 周）", level=2)
p = doc.add_paragraph()
run = p.add_run("仅在验证成功后启动")
run.bold = True
phase4 = [
    "基于真实需求设计系统架构（不是猜测）",
    "优先开发核心瓶颈功能",
    "逐步从工具组合迁移到自建平台",
    "目标：支持 100+ 客户的全自动化运营",
]
for item in phase4:
    doc.add_paragraph(item, style="List Bullet")

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 成本对比总结
# ══════════════════════════════════════════════════════════
doc.add_heading("五、成本对比总结", level=1)

doc.add_heading("场景 1：验证成功 → 最终自建", level=2)
add_table(
    ["阶段", "时间", "投入", "产出"],
    [
        ["工具组合验证", "第 1-12 周", "~$1,500", "验证数据 + 5-20 个客户"],
        ["自建平台开发", "第 13-24 周", "~$8,000-15,000", "定制平台"],
        ["合计", "24 周", "~$10,000-17,000", "有数据支撑的平台"],
    ],
)

doc.add_heading("场景 2：验证失败 → 止损", level=2)
add_table(
    ["阶段", "时间", "投入", "产出"],
    [
        ["工具组合验证", "第 1-12 周", "~$1,500", "宝贵的市场认知"],
        ["止损退出", "—", "$0", "避免了 $10,000+ 的无效开发"],
        ["合计", "12 周", "~$1,500", "低成本试错"],
    ],
)

doc.add_heading("场景 3：直接自建（不推荐）", level=2)
add_table(
    ["阶段", "时间", "投入", "产出"],
    [
        ["开发平台", "第 1-12 周", "~$10,000-30,000", "未经验证的平台"],
        ["找客户验证", "第 13-20 周", "~$1,000", "如果失败，前期开发全部浪费"],
        ["合计", "20 周", "~$11,000-31,000", "高风险"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 最终建议
# ══════════════════════════════════════════════════════════
doc.add_heading("六、最终建议", level=1)

p = doc.add_paragraph()
run = p.add_run("推荐路径：方案 A（工具组合）→ 验证 → 方案 B（自建平台）")
run.bold = True
run.font.size = Pt(13)

doc.add_paragraph()

reasons = [
    ("1. 速度优先", "1-2 周上线 vs 8-12 周开发。先占市场，先有客户。"),
    ("2. 风险可控", "最大损失 ~$1,500 vs 最大损失 $10,000-30,000。"),
    ("3. 数据驱动", "基于真实运营数据决定是否自建、建什么、怎么建。"),
    ("4. 技术降险", "AI 视频生成技术每 3 个月都在进步，等一等可能成本更低、质量更好。"),
    ("5. 聚焦核心", "先验证「AI 营销自动化是否有人买单」，而不是「我能不能建一个平台」。"),
]

for title_text, desc in reasons:
    p = doc.add_paragraph()
    run = p.add_run(title_text + "  ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("-- SocialChef Build vs Buy 决策分析 -- 2026.04.13 --")
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(10)

# Save
output = os.path.expanduser("~/Desktop/SocialChef_Build_vs_Buy.docx")
doc.save(output)
print(f"Saved to: {output}")
